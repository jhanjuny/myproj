"""
DOCX processor using python-docx.
Extracts paragraphs (with heading level), tables, and embedded images.
"""
from __future__ import annotations

import io
from pathlib import Path

from src.config import MAX_IMAGE_BYTES, MAX_IMAGES_PER_FILE
from src.processors.base import BaseProcessor, ImageBlock, ProcessedContent, ProcessorError


class DocxProcessor(BaseProcessor):
    supported_extensions = (".docx", ".doc")

    def process(self, path: Path) -> ProcessedContent:
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            raise ProcessorError("python-docx is not installed (pip install python-docx)")

        try:
            doc = Document(str(path))
        except Exception as e:
            raise ProcessorError(f"Cannot open DOCX: {e}")

        text_parts: list[str] = []
        images: list[ImageBlock] = []
        image_count = 0

        # ── Paragraphs ────────────────────────────────────────────────────────
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            txt = para.text.strip()
            if not txt:
                continue
            # Convert Word heading styles to markdown-like markers
            if style.startswith("Heading 1"):
                text_parts.append(f"# {txt}")
            elif style.startswith("Heading 2"):
                text_parts.append(f"## {txt}")
            elif style.startswith("Heading 3"):
                text_parts.append(f"### {txt}")
            else:
                text_parts.append(txt)

        # ── Tables ────────────────────────────────────────────────────────────
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                text_parts.append("\n[Table]\n" + "\n".join(rows))

        # ── Embedded images ───────────────────────────────────────────────────
        for rel in doc.part.rels.values():
            if image_count >= MAX_IMAGES_PER_FILE:
                break
            if "image" in rel.reltype:
                try:
                    raw = rel.target_part.blob
                    ct = rel.target_part.content_type  # e.g. "image/png"
                    if len(raw) > MAX_IMAGE_BYTES:
                        raw = _downscale(raw, MAX_IMAGE_BYTES)
                        ct = "image/jpeg"
                    images.append(ImageBlock(
                        data=raw,
                        media_type=ct,
                        caption="Embedded image",
                        page_or_slide=0,
                    ))
                    image_count += 1
                except Exception:
                    continue

        metadata = {
            "title": path.stem,
            "paragraph_count": len(doc.paragraphs),
        }

        content = ProcessedContent(
            text="\n\n".join(text_parts),
            images=images,
            metadata=metadata,
            source_path=path,
            file_type="docx",
        )

        if content.is_empty():
            raise ProcessorError(f"Could not extract any content from {path.name}")

        return content


def _downscale(raw: bytes, max_bytes: int) -> bytes:
    from PIL import Image

    img = Image.open(io.BytesIO(raw)).convert("RGB")
    for q in (75, 55, 35):
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=q)
        data = buf.getvalue()
        if len(data) <= max_bytes:
            return data
    img.thumbnail((800, 800))
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=50)
    return buf.getvalue()
